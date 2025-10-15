from docx import Document
import pandas as pd
import os

def generate_document(data):
    doc = Document()
    main_heading = f"{data['LID-NR']} -- {data['Gene'].upper()}"
    p = doc.add_paragraph()
    run = p.add_run(main_heading)
    run.bold = True
    run.font.size = 120000  # adjust size as needed

    separator = "=============================================="
    spaces = ' ' * 10

    doc.add_paragraph(separator)
    p = doc.add_paragraph()
    run = p.add_run(f"Genetisk Variant:")
    run.bold = True
    run.font.size = 100000  # adjust size as needed
    variant_info = process_variant_info(data, spaces)
    doc.add_paragraph(spaces + variant_info)

    doc.add_paragraph(separator)
    p = doc.add_paragraph()
    run = p.add_run(f"Bedömning och Databasreferenser:")
    run.bold = True
    run.font.size = 100000  # adjust size as needed
    doc.add_paragraph(spaces + f"Bedömning enligt ACMG-kriterierna: {data['ACMG criteria assessment']}.")
    doc.add_paragraph(spaces + f"Tidigare rapporter i ClinVar och hemofilidatabaserna: {data['ClinVar and hemophilia database reports']}.")

    if 'Inversions' in data and data['Inversions']:
        doc.add_paragraph(separator)
        p = doc.add_paragraph()
        run = p.add_run("Utredning av Inversionsvarianter:")
        run.bold = True
        run.font.size = 100000  # adjust size as needed
        if data['Inversions'] == 'ja':
            doc.add_paragraph(spaces + f"Resultat av inversionerna i intron 1 och 22: {data.get('Inversion result', 'Inga resultat tillgängliga')}.")
        elif data['Inversions'] == 'nej':
            doc.add_paragraph(spaces + f"Utredning av inversionsvarianter i intron 1 och 22 utfördes inte på grund av: {data.get('Inversion reason', 'Ingen anledning angiven')}.")

    # CSV data förberedelse
    csv_row = prepare_csv_row(data['LID-NR'], data)
    csv_data_rows = [csv_row]

    doc.add_paragraph(separator)
    p = doc.add_paragraph()
    run = p.add_run("Proband:")
    run.bold = True
    run.font.size = 100000  # adjust size as needed
    if data['Proband'] == "Proband inte känt":
        doc.add_paragraph(spaces + "Proband inte känt")
    else:
        doc.add_paragraph(spaces + f"Pnr: {data['Proband']}")
        doc.add_paragraph(spaces + f"Genotyp: {data['Genotype']}")
        doc.add_paragraph(spaces + f"Fenotyp: {data['Phenotype']}")

    doc.add_paragraph(separator)
    p = doc.add_paragraph()
    run = p.add_run("Genomisk Analys:")
    run.bold = True
    run.font.size = 100000  # adjust size as needed
    if data['Sequencing method'] == 'mps':
        doc.add_paragraph(
            spaces + "Alla kodande regioner med flankerande icke-kodande sekvenser har analyserats. Sekvensdata har mappats mot referenssekvens (GRCh37/hg19). "
                     "Analysen innefattar endast exon och exon-intron-gränser och därför ingår inte promotor-, intron- och icke-kodande-regioner i analysen."
        )
    else:
        doc.add_paragraph(spaces + f"Exon {data['Exon']} av {data['Gene'].upper()} har analyserats med Sangersekvensering.")

    doc.add_paragraph(separator)
    p = doc.add_paragraph()
    run = p.add_run("Referenser:")
    run.bold = True
    run.font.size = 100000  # adjust size as needed
    doc.add_paragraph(spaces + "1. Genome Reference Consortium Human Build (2022), Vol. 37.")
    doc.add_paragraph(spaces + "2. Clinical genomics (2022), Stockholm, Sverige. Tillgängligt vid: https://www.scilifelab.se/facilities/clinical-genomics-stockholm/")
    doc.add_paragraph(spaces + "3. Scout (2022), Tillgängligt vid: https://github.com/Clinical-Genomics/scout")
    doc.add_paragraph(spaces + "4. Landrum MJ, Lee JM, Benson M, Brown GR, Chao C, Chitipiralla S, Gu B, Hart J, Hoffman D, Jang W, Karapetyan K, Katz K, Liu C, Maddipatla Z, Malheiro A, McDaniel K, Ovetsky M, Riley G, Zhou G, Holmes JB, Kattman BL, Maglott DR. ClinVar: improving access to variant interpretations and supporting evidence. Nucleic Acids Res, 2018 Jan 4. PubMed PMID: 29165669")
    doc.add_paragraph(spaces + "5. EAHAD Coagulation Factor Variant Databases (2022): https://dbs.eahad.org")
    doc.add_paragraph(spaces + "6. Richards, Sue et al. Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology. Genetics in medicine: official journal of the American College of Medical Genetics vol. 17,5 (2015): 405-24. doi:10.1038/gim.2015.30")

    doc.add_paragraph(separator)
    p = doc.add_paragraph()
    run = p.add_run("Avsändare:")
    run.bold = True
    run.font.size = 100000  # adjust size as needed
    doc.add_paragraph(spaces + "Professor, Överläkare: Jovan Antovic")
    doc.add_paragraph(spaces + "Sjukhuskemist: August Lundholm")
    doc.add_paragraph(spaces + "Biomedicinsk analytiker: Somia Echehli")

    # Spara dokument
    save_path = os.path.join('H:', 'DNA_Sekvenseringsresultat', 'Remissvar Koagulation', 'Väntande på att svaras ut', f"{data['LID-NR']} ({data['Gene'].upper()})", 'genetisk_rapport.docx')
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    print(f"Dokumentet har skapats och sparats som '{save_path}'")

    # Uppdatera CSV-fil om det finns data
    if csv_data_rows:
        csv_path = 'H:\\DNA_Sekvenseringsresultat\\Remissvar Koagulation\\intressanta_varianter.csv'
        update_csv_file(csv_path, pd.DataFrame(csv_data_rows))

def process_variant_info(variant, spaces):
    """Bearbetar och returnerar textinformation för en genetisk variant."""
    info = f"{variant['Gene'].upper()} ({variant['Transcript']}): c. {variant['Nucleotide change']} p. {variant['Protein change']} {variant['Zygosity']}."
    info += f"\n{spaces}Det är troligt att varianten orsakar {variant['Disease']}."
    info += f"\n{spaces}Nedärvning: {variant['Inheritance']}."
    return info

def prepare_csv_row(lid_nr, variant):
    """Förbereder en rad med data för att lägga till i CSV-filen."""
    return {
        'LID-NR': lid_nr,
        'Gene': variant['Gene'].upper(),
        'Nucleotide change': variant['Nucleotide change'],
               'Protein change': variant['Protein change'],
        'Zygosity': variant['Zygosity'],
        'Inheritance': variant['Inheritance'],
        'ACMG criteria assessment': variant['ACMG criteria assessment'],
        'ClinVar and hemophilia database reports': variant['ClinVar and hemophilia database reports']
    }

def update_csv_file(csv_path, new_data):
    """Uppdaterar en CSV-fil med ny data."""
    try:
        if os.path.exists(csv_path):
            new_data.to_csv(csv_path, mode='a', header=False, index=False)
        else:
            new_data.to_csv(csv_path, index=False)
        print("CSV-filen har uppdaterats.")
    except Exception as e:
        print(f"Ett fel uppstod när CSV-filen skulle uppdateras: {e}")

