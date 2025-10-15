import unittest
import os
from generate_document import generate_document
from docx import Document

class TestGenerateDocument(unittest.TestCase):
    def setUp(self):
        # Skapa en testdata med två varianter
        self.test_data = {
            'LID-NR': 'test123',
            'Proband': 'Proband1',
            'Genotype': 'Genotype1',
            'Phenotype': 'Phenotype1',
            'Sequencing method': 'mps',
            'Exon': 'Exon1',
            'variants': [
                {
                    'Gene': 'f8',
                    'Transcript': 'NM_000132',
                    'Disease': 'Hemofili A',
                    'Nucleotide change': 'c.1234A>G',
                    'Protein change': 'p.Tyr412Cys',
                    'Zygosity': 'Hemizygot',
                    'Inheritance': 'Autosomalt dominant',
                    'ACMG criteria assessment': 'Patogen',
                    'ClinVar and hemophilia database reports': 'ClinVar data',
                    'Further studies': 'ja'
                },
                {
                    'Gene': 'vwf',
                    'Transcript': 'NM_000552',
                    'Disease': 'von Willebrands sjukdom',
                    'Nucleotide change': 'c.5678G>A',
                    'Protein change': 'p.Ala123Thr',
                    'Zygosity': 'Heterozygot',
                    'Inheritance': 'Autosomalt recessiv',
                    'ACMG criteria assessment': 'Troligen patogen',
                    'ClinVar and hemophilia database reports': 'ClinVar report',
                    'Further studies': 'nej'
                }
            ]
        }
        self.output_path = 'test_output'
        os.makedirs(self.output_path, exist_ok=True)
    
    def tearDown(self):
        # Rensa testdokumentet efter testet
        file_path = os.path.join(self.output_path, 'genetisk_rapport.docx')
        if os.path.exists(file_path):
            os.remove(file_path)
    
    def test_generate_document(self):
        generate_document(self.test_data, self.output_path)
        file_path = os.path.join(self.output_path, 'genetisk_rapport.docx')
        self.assertTrue(os.path.exists(file_path))

        # Kontrollera innehållet i det genererade dokumentet
        doc = Document(file_path)
        paragraphs = [p.text.replace(' ', '').replace('\n', '') for p in doc.paragraphs]

        # Kontrollera första varianten
        self.assertIn('test123--F8', paragraphs)
        self.assertIn('GenetiskVariant1:', paragraphs)
        self.assertIn('F8(NM_000132):c.c.1234A>Gp.p.Tyr412CysHemizygot.DetärtroligtattvariantenorsakarHemofiliA.Nedärvning:Autosomaltdominant.', paragraphs)
        self.assertIn('BedömningenligtACMG-kriterierna:Patogen.', paragraphs)
        self.assertIn('TidigarerapporteriClinVarochhemofilidatabaserna:ClinVardata.', paragraphs)

        # Kontrollera andra varianten
        self.assertIn('GenetiskVariant2:', paragraphs)
        self.assertIn('VWF(NM_000552):c.c.5678G>Ap.p.Ala123ThrHeterozygot.DetärtroligtattvariantenorsakarvonWillebrandssjukdom.Nedärvning:Autosomaltrecessiv.', paragraphs)
        self.assertIn('BedömningenligtACMG-kriterierna:Troligenpatogen.', paragraphs)
        self.assertIn('TidigarerapporteriClinVarochhemofilidatabaserna:ClinVarreport.', paragraphs)

if __name__ == '__main__':
    unittest.main()
