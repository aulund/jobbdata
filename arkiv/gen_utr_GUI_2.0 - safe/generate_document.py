import os
from docx import Document
from docx.shared import Pt
from datetime import datetime

def generate_document(data, output_dir="output"):
    lid_nr = data['LID-NR']
    gene_name = data['variants'][0]['Gene'].upper()
    filename = f"{lid_nr}_{gene_name}.docx"
    filepath = os.path.join(output_dir, filename)

    document = Document()
    document.add_heading(f"{lid_nr} -- {gene_name}", level=1)
    document.add_paragraph("==============================")

    for idx, variant in enumerate(data['variants'], start=1):
        gene = variant['Gene'].upper()
        transcript = variant['Transcript']
        disease = variant['Disease']
        nucleotide_change = variant['Nucleotide change']
        protein_change = variant['Protein change']
        zygosity = variant['Zygosity']
        inheritance = variant['Inheritance']
        acmg = variant['ACMG criteria assessment']
        clinvar = variant['ClinVar and hemophilia database reports']
        further_studies = variant['Further studies']

        variant_heading = f"Genetisk Variant {idx}:"
        variant_info = (f"{gene} ({transcript}): c. {nucleotide_change} p. {protein_change} {zygosity}.\n"
                        f"          Det är troligt att varianten orsakar {disease}.\n"
                        f"          Nedärvning: {inheritance}.")
        
        document.add_paragraph(variant_heading)
        document.add_paragraph("==============================")
        document.add_paragraph(variant_info)
        document.add_paragraph("==============================")
        document.add_paragraph("Bedömning och Databasreferenser för Variant {idx}:")
        document.add_paragraph(f"          Bedömning enligt ACMG-kriterierna: {acmg}.")
        document.add_paragraph(f"          Tidigare rapporter i ClinVar och hemofilidatabaserna: {clinvar}.")
        document.add_paragraph("==============================")

    proband_info = ("Proband:",
                    f"          Pnr: {data['Proband']}",
                    f"          Genotyp: {data['Genotype']}",
                    f"          Fenotyp: {data['Phenotype']}")
    
    sequencing_method = data['Sequencing method']
    exon_info = data['Exon'] if sequencing_method == "Sanger" else "Alla kodande regioner med flankerande icke-kodande sekvenser har analyserats. Sekvensdata har mappats mot referenssekvens (GRCh37/hg19). Analysen innefattar endast exon och exon-intron-gränser och därför ingår inte promotor-, intron- och icke-kodande-regioner i analysen."

    genomic_analysis = ("Genomisk Analys:",
                        f"          {exon_info}")

    references = ("Referenser:",
                  "          1. Genome Reference Consortium Human Build (2022), Vol. 37.",
                  "          2. Clinical genomics (2022), Stockholm, Sverige. Tillgängligt vid: https://www.scilifelab.se/facilities/clinical-genomics-stockholm/",
                  "          3. Scout (2022), Tillgängligt vid: https://github.com/Clinical-Genomics/scout",
                  "          4. Landrum MJ, Lee JM, Benson M, Brown GR, Chao C, Chitipiralla S, Gu B, Hart J, Hoffman D, Jang W, Karapetyan K, Katz K, Liu C, Maddipatla Z, Malheiro A, McDaniel K, Ovetsky M, Riley G, Zhou G, Holmes JB, Kattman BL, Maglott DR. ClinVar: improving access to variant interpretations and supporting evidence. Nucleic Acids Res, 2018 Jan 4. PubMed PMID: 29165669",
                  "          5. EAHAD Coagulation Factor Variant Databases (2022): https://dbs.eahad.org",
                  "          6. Richards, Sue et al. Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology. Genetics in medicine: official journal of the American College of Medical Genetics vol. 17,5 (2015): 405-24. doi:10.1038/gim.2015.30")

    sender_info = ("Avsändare:",
                   "          Professor, Överläkare",
                   "          Namn",
                   "          Klinik",
                   "          Sjukhus",
                   "          Kontaktinformation")

    for paragraph in proband_info + ("==============================",) + genomic_analysis + ("==============================",) + references + ("==============================",) + sender_info:
        document.add_paragraph(paragraph)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    document.save(filepath)
    print(f"Dokumentet har skapats och sparats som '{filepath}'")
