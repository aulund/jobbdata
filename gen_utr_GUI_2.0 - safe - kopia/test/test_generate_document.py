import os
import unittest
from unittest.mock import patch, MagicMock
from docx import Document
from generate_document import generate_document  # This should match the actual file name

class TestGenerateDocument(unittest.TestCase):
    
    @patch('generate_document.Document')
    def test_generate_document(self, MockDocument):
        # Set up mock document
        mock_doc = MagicMock()
        MockDocument.return_value = mock_doc
        
        # Ensure the mock has the expected methods
        mock_doc.add_heading = MagicMock()
        mock_doc.add_paragraph = MagicMock()
        mock_doc.save = MagicMock()
        
        # Example data to test with
        data = {
            "LID-NR": "1234",
            "Proband": "Proband Name",
            "Genotype": "Genotype Example",
            "Phenotype": "Phenotype Example",
            "Sequencing method": "MPS",
            "Exon": "2",
            "variants": [
                {
                    "Gene": "F8",
                    "Transcript": "NM_000132",
                    "Disease": "Hemophilia A",
                    "Nucleotide change": "c.1234G>A",
                    "Protein change": "p.Val412Met",
                    "Zygosity": "Heterozygot",
                    "Inheritance": "X-linked",
                    "ACMG criteria assessment": "Benign",
                    "ClinVar and hemophilia database reports": "ClinVar Report Example",
                    "Further studies": "No"
                }
            ]
        }
        
        # Mock the output directory
        output_directory = r"./test_output"

        # Call the function under test
        generate_document(data, output_directory)
        
        # Print all calls made to add_paragraph
        print("Calls to add_paragraph:")
        for call in mock_doc.add_paragraph.call_args_list:
            print(call)

        # Assert that a document was created
        MockDocument.assert_called_once()
        mock_doc.save.assert_called_once_with(os.path.join(output_directory, "1234_F8.docx"))
        
        # Check that the mock document has had the correct content added
        mock_doc.add_heading.assert_any_call("1234 -- F8", level=1)
        
        # Attempt to match the expected call
        mock_doc.add_paragraph.assert_any_call(
    'F8 (NM_000132): c.1234G>A p.Val412Met Heterozygot.\nDet är troligt att varianten orsakar Hemophilia A.\nNedärvning: X-linked.'
)

        
        mock_doc.add_paragraph.assert_any_call("Proband:")
        mock_doc.add_paragraph.assert_any_call("Proband Name\nGenotyp: Genotype Example\nFenotyp: Phenotype Example")

if __name__ == '__main__':
    unittest.main()
