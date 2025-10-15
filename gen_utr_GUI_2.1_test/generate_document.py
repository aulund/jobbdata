import os
from docx import Document

def generate_document(data, output_path):
    try:
        print("Starting report generation...")

        if not data:
            print("Ingen data att generera dokument från.")
            return

        # Debug: Print the data received
        print("Data received:", data)

        doc = Document()
        main_heading = f"{data['LID-NR']} -- {data['variants'][0]['Gene'].upper()}"
        doc.add_heading(main_heading, level=1)

        doc.add_paragraph("------------------------------")
        for i, variant in enumerate(data['variants'], 1):
            # Ensure Transcript and Disease fields are correctly assigned
            variant['Transcript'] = variant.get('Transcript', data.get('Transcript', 'N/A'))
            variant['Disease'] = variant.get('Disease', data.get('Disease', 'N/A'))
            
            variant_info = process_variant_info(variant)
            doc.add_paragraph(f"Genetisk Variant {i}:")
            doc.add_paragraph(variant_info)

            # Add ACMG and database reports if present
            assessment_info = process_assessment_info(variant)
            doc.add_paragraph(assessment_info)
            
            doc.add_paragraph("------------------------------")

        proband_info = process_proband_info(data)
        doc.add_paragraph("Proband:")
        doc.add_paragraph(proband_info)
        doc.add_paragraph("------------------------------")

        genomic_analysis_info = process_genomic_analysis_info(data)
        doc.add_paragraph("Genomisk Analys:")
        doc.add_paragraph(genomic_analysis_info)
        doc.add_paragraph("------------------------------")

        #references = get_references()
        #doc.add_paragraph("Referenser:")
        #for ref in references:
        #    doc.add_paragraph(ref)
        #doc.add_paragraph("------------------------------")

        sender_info = get_sender_info()
        doc.add_paragraph("Avsändare:")
        doc.add_paragraph(sender_info)
        doc.add_paragraph("------------------------------")

        filename = f"{data['LID-NR']}_{data['variants'][0]['Gene'].upper()}.docx"
        output_file = os.path.join(output_path, filename)
        doc.save(output_file)
        print(f"Dokumentet har skapats och sparats som '{output_file}'")
    except Exception as e:
        print(f"An error occurred during report generation: {e}")

def process_variant_info(variant):
    gene = variant.get('Gene', 'N/A')
    transcript = variant.get('Transcript', 'N/A')
    nucleotide_change = variant.get('Nucleotide change', 'N/A').strip()
    protein_change = variant.get('Protein change', 'N/A').strip()
    zygosity = variant.get('Zygosity', 'N/A')
    inheritance = variant.get('Inheritance', 'N/A')
    disease = variant.get('Disease', 'N/A')
    
    # Clean up the unnecessary repetitions and correct formatting
    info = (f"{gene} ({transcript}): c.{nucleotide_change} p.{protein_change} {zygosity}.\n"
            f"Det är troligt att varianten orsakar {disease}.\n"
            f"Nedärvning: {inheritance}.")
    return info

def process_assessment_info(variant):
    acmg = variant.get('ACMG criteria assessment', 'N/A')
    clinvar = variant.get('ClinVar and hemophilia database reports', 'N/A')
    
    info = (f"Bedömning enligt ACMG-kriterierna: {acmg}.\n"
            f"Sammanfattning och utlåtande: {clinvar}.")
    return info

def process_proband_info(data):
    proband = data.get('Proband', 'N/A')
    genotype = data.get('Genotype', 'N/A')
    phenotype = data.get('Phenotype', 'N/A')

    info = f"{proband}\n"
    if proband != 'N/A':  # If proband has a meaningful value, add Genotype and Phenotype
        info += (f"Genotyp: {genotype}\n"
                 f"Fenotyp: {phenotype}")
    return info

def process_genomic_analysis_info(data):
    sequencing_method = data.get('Sequencing method', 'N/A').strip().lower()
    exon = data.get('Exon', 'N/A')

    if sequencing_method == 'sanger':
        info = f"Exon {exon} av {data['variants'][0]['Gene']} har analyserats med Sangersekvensering."
    else:  # Handle MPS or any other method
        info = ("Alla kodande regioner med flankerande icke-kodande sekvenser har analyserats med MPS. "
                "Sekvensdata har mappats mot referenssekvens (GRCh37/hg19). "
                "Analysen innefattar endast exon och exon-intron-gränser och därför ingår inte "
                "promotor-, intron- och icke-kodande-regioner i analysen.")
    
    return info


def get_references():
    references = [
     #   "1. Genome Reference Consortium Human Build (2022), Vol. 37.",
      #  "2. Clinical genomics (2022), Stockholm, Sverige. Tillgängligt vid: https://www.scilifelab.se/facilities/clinical-genomics-stockholm/",
       # "3. Scout (2022), Tillgängligt vid: https://github.com/Clinical-Genomics/scout",
        #"4. Landrum MJ, Lee JM, Benson M, Brown GR, Chao C, Chitipiralla S, Gu B, Hart J, Hoffman D, Jang W, Karapetyan K, Katz K, Liu C, Maddipatla Z, Malheiro A, McDaniel K, Ovetsky M, Riley G, Zhou G, Holmes JB, Kattman BL, Maglott DR. ClinVar: improving access to variant interpretations and supporting evidence. Nucleic Acids Res, 2018 Jan 4. PubMed PMID: 29165669",
        #"5. EAHAD Coagulation Factor Variant Databases (2022): https://dbs.eahad.org",
        #"6. Richards, Sue et al. Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology. Genetics in medicine: official journal of the American College of Medical Genetics vol. 17,5 (2015): 405-24. doi:10.1038/gim.2015.30"
    ]
    return references

def get_sender_info():
    sender_info = ("MVH August Lundholm, Molekylärbiolog")
    return sender_info
