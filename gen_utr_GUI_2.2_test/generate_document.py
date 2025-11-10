import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config

logger = logging.getLogger(__name__)

# Template file name (can be customized)
TEMPLATE_FILE = "variant_report_template.docx"

# Laboratory information
LAB_NAME = "KLINISK UNIVERSITETSLABORATORIET"
LAB_TITLE = "Svarsrapport"
LAB_ACCREDITATION = "Ackrediterat laboratorium enligt ISO 15189"
LAB_ADDRESS = "Karolinska Universitetssjukhuset\nKlinisk kemi\n171 76 Stockholm"
LAB_CONTACT = """Karolinska Universitetslaboratoriet (kod: XXXXXX)
171 76 Stockholm
Tel: XX-XXX XX XX | Fax: XX-XXX XX XX
E-post: kontakt@karolinska.se"""


def set_arial_font(run, size=11, bold=False):
    """Set Arial font for a run."""
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heavy_border_box(paragraph):
    """
    Add heavy black borders around a paragraph to create a box.
    
    Args:
        paragraph: The paragraph to add borders to
    """
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Add borders to all sides with heavier weight
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')  # Heavy border (3pt)
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '000000')  # Black
        pBdr.append(border)
    
    pPr.append(pBdr)


def create_header_section(doc, data, page_num=1, total_pages=2):
    """
    Create the header section matching clinical laboratory format.
    Uses data from actual input fields.
    
    Args:
        doc: Document object
        data: Dictionary with patient/sample data
        page_num: Current page number
        total_pages: Total number of pages
    """
    # Title and page number row
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    
    # Logo cell (left) - placeholder for when template is not used
    logo_cell = header_table.rows[0].cells[0]
    p = logo_cell.paragraphs[0]
    run = p.add_run(LAB_NAME)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Title cell (center)
    title_cell = header_table.rows[0].cells[1]
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(LAB_TITLE)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Page number cell (right)
    page_cell = header_table.rows[0].cells[2]
    p = page_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Sida {page_num} av {total_pages}")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    # Set column widths
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(1.0)
    
    # Accreditation line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(LAB_ACCREDITATION)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing
    
    # Patient/Sample info and Test details side by side
    info_table = doc.add_table(rows=1, cols=2)
    
    # Left side: Patient/Sample info (use actual data or defaults)
    left_cell = info_table.rows[0].cells[0]
    left_cell.width = Inches(3.5)
    
    p = left_cell.add_paragraph()
    run = p.add_run("Svar till:")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Use Category for facility name if available
    facility = "Karolinska Universitetssjukhuset, DNA-laboratoriet"
    category = data.get('Category', '')
    if category:
        facility = f"Karolinska H, {category}"
    
    p = left_cell.add_paragraph()
    run = p.add_run(facility)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    p = left_cell.add_paragraph()
    run = p.add_run(LAB_ADDRESS)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    # Right side: Test details (use actual data)
    right_cell = info_table.rows[0].cells[1]
    right_cell.width = Inches(3.0)
    
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Slutsvar medicinskt granskat")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    response_date = datetime.now().strftime("%Y-%m-%d")
    run = p.add_run(f"Svarsdatum: {response_date}")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lab_id = data.get('LID-NR', 'N/A')
    run = p.add_run(f"LabId: {lab_id}")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    
    # Note: Arrival date and Remiss ID are not collected in the GUI,
    # so we don't add them unless they're provided
    
    doc.add_paragraph()  # Spacing


def create_footer_section(doc):
    """
    Create the footer section with biobank info and laboratory contact.
    
    Args:
        doc: Document object
    """
    doc.add_paragraph()  # Spacing
    
    # Biobank information
    p = doc.add_paragraph()
    run = p.add_run("Biobankinformation: Prov kan sparas för kvalitetssäkring enligt Biobankslagen.")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.underline = True
    
    # Version and date (right aligned)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    version_date = datetime.now().strftime("%Y-%m-%d")
    run = p.add_run(f"Version 2.6 | {version_date}")
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    
    doc.add_paragraph()  # Spacing
    
    # Laboratory contact information
    p = doc.add_paragraph()
    run = p.add_run(LAB_CONTACT)
    run.font.name = 'Arial'
    run.font.size = Pt(9)


def create_boxed_section(doc, title, content_paragraphs):
    """
    Create a content section with heavy black borders.
    
    Args:
        doc: Document object
        title: Section title
        content_paragraphs: List of content strings or tuples (text, bold)
    """
    # Create a table for the boxed section
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # Apply heavy borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')  # Heavy border (3pt)
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    cell = table.rows[0].cells[0]
    
    # Add title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    # Add content paragraphs
    for content in content_paragraphs:
        p = cell.add_paragraph()
        
        if isinstance(content, tuple):
            text, bold = content
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = bold
        else:
            run = p.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        
        p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()  # Spacing after box


def add_horizontal_line(paragraph):
    """
    Add a horizontal line (border) to a paragraph.
    
    Args:
        paragraph: The paragraph to add the border to
    """
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Border size
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')  # Professional blue color
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_heading(paragraph, text, level=1, color=None):
    """
    Apply professional styling to a heading.
    
    Args:
        paragraph: The paragraph to style
        text: The text content
        level: Heading level (1-3)
        color: Optional RGB color tuple (r, g, b)
    """
    run = paragraph.add_run(text)
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(68, 114, 196) if not color else RGBColor(*color)  # Professional blue
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(68, 114, 196) if not color else RGBColor(*color)
    else:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(89, 89, 89) if not color else RGBColor(*color)  # Dark gray
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_section_heading(doc, text, level=2):
    """
    Add a professional section heading with styling.
    
    Args:
        doc: Document object
        text: Heading text
        level: Heading level
        
    Returns:
        The created paragraph
    """
    paragraph = doc.add_paragraph()
    style_heading(paragraph, text, level)
    return paragraph


def add_styled_paragraph(doc, text, bold=False, indent=False):
    """
    Add a styled paragraph with consistent formatting.
    
    Args:
        doc: Document object
        text: Paragraph text
        bold: Whether to make text bold
        indent: Whether to indent the paragraph
        
    Returns:
        The created paragraph
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    if bold:
        run.font.bold = True
    
    if indent:
        paragraph.paragraph_format.left_indent = Inches(0.25)
    
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_spacer(doc, size=12):
    """
    Add vertical spacing between sections.
    
    Args:
        doc: Document object
        size: Space size in points
    """
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(size)


def add_variant_table(doc, variant, variant_number):
    """
    Add a professionally formatted table for variant information with proper borders.
    
    Args:
        doc: Document object
        variant: Dictionary containing variant details
        variant_number: The number/index of this variant
    """
    # Add variant heading
    add_section_heading(doc, f"Genetisk Variant {variant_number}", level=2)
    add_spacer(doc, 6)
    
    # Create table with 2 columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'  # Use Table Grid for visible borders
    
    # Apply table borders explicitly
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    def set_table_borders(table):
        """Set consistent borders for the entire table"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Create table borders element
        tblBorders = OxmlElement('w:tblBorders')
        
        # Define border settings for all sides
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border width in eighths of a point
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black borders
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    set_table_borders(table)
    
    # Helper function to add a row
    def add_row(label, value):
        row = table.add_row()
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        
        # Style label cell
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = RGBColor(68, 114, 196)
        label_para.paragraph_format.space_before = Pt(3)
        label_para.paragraph_format.space_after = Pt(3)
        
        # Style value cell
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(str(value))
        value_run.font.size = Pt(11)
        value_run.font.name = 'Calibri'
        value_para.paragraph_format.space_before = Pt(3)
        value_para.paragraph_format.space_after = Pt(3)
    
    # Add variant details
    gene = variant.get('Gene', config.DEFAULT_UNKNOWN_VALUE)
    transcript = variant.get('Transcript', config.DEFAULT_UNKNOWN_VALUE)
    nucleotide_change = variant.get('Nucleotide change', config.DEFAULT_UNKNOWN_VALUE).strip()
    protein_change = variant.get('Protein change', config.DEFAULT_UNKNOWN_VALUE).strip()
    zygosity = variant.get('Zygosity', config.DEFAULT_UNKNOWN_VALUE)
    inheritance = variant.get('Inheritance', config.DEFAULT_UNKNOWN_VALUE)
    disease = variant.get('Disease', config.DEFAULT_UNKNOWN_VALUE)
    acmg = variant.get('ACMG criteria assessment', config.DEFAULT_UNKNOWN_VALUE)
    clinvar = variant.get('ClinVar and hemophilia database reports', config.DEFAULT_UNKNOWN_VALUE)
    
    add_row("Gen", f"{gene}")
    add_row("Transkript", f"{transcript}")
    add_row("Nukleotidförändring", f"c.{nucleotide_change}")
    add_row("Proteinförändring", f"p.{protein_change}")
    add_row("Zygositet", f"{zygosity}")
    add_row("Nedärvning", f"{inheritance}")
    add_row("Sjukdom", f"{disease}")
    add_row("ACMG-bedömning", f"{acmg}")
    
    # Set column widths with better proportions
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    
    add_spacer(doc, 6)
    
    # Add clinical significance paragraph
    if clinvar != config.DEFAULT_UNKNOWN_VALUE and clinvar.strip():
        add_styled_paragraph(doc, "Sammanfattning och utlåtande:", bold=True)
        add_styled_paragraph(doc, clinvar, indent=True)
    
    add_spacer(doc, 12)


def generate_document(data, output_path):
    """
    Generate a clinical laboratory report document in Swedish following ISO 15189 format.
    
    Args:
        data: Dictionary containing all collected variant and patient information
        output_path: Directory path where the document should be saved
        
    Returns:
        str: Path to the generated document, or None if generation failed
    """
    try:
        logger.info("Starting clinical report generation...")

        if not data:
            logger.error("No data provided for document generation")
            return None
        
        # Validate output path
        if not output_path or '..' in output_path:
            logger.error(f"Invalid output path: {output_path}")
            return None

        # Check if this is a normal finding
        if data.get("Normalfynd", False):
            return generate_normalfinding_document(data, output_path)

        # Validate required data
        if not data.get("variants"):
            logger.error("No variants provided in data")
            return None
            
        if not data.get("LID-NR"):
            logger.error("No LID-NR provided in data")
            return None

        # Load template with logos and formatting
        template_path = os.path.join(os.path.dirname(__file__), TEMPLATE_FILE)
        if os.path.exists(template_path):
            doc = Document(template_path)
            logger.info(f"Using template with logos: {template_path}")
            
            # Fill in LID-NR in the template table
            if doc.tables and len(doc.tables) > 0:
                template_table = doc.tables[0]
                for row in template_table.rows:
                    for cell in row.cells:
                        if "LID" in cell.text:
                            # Clear and set LID value
                            cell.text = f"LID: {data['LID-NR']}"
                            # Apply formatting from template's style
                            if cell.paragraphs:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.font.size = Pt(11)
                                        run.font.bold = True
        else:
            doc = Document()
            logger.warning("Template not found, creating blank document")
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
        
        # === PAGE 1: HEADER ===
        create_header_section(doc, data, page_num=1, total_pages=2)
        
        # === PAGE 1: SUMMARY AND FINDINGS BOX ===
        summary_content = []
        
        # Test description
        test_type = "DNA-analys av genetiska varianter"
        summary_content.append(("Testtyp:", True))
        summary_content.append(test_type)
        summary_content.append("")  # Empty line
        
        # Clinical findings
        summary_content.append(("Kliniska fynd:", True))
        
        # Determine finding status
        finding_status = "PATOGEN VARIANT PÅVISAD"
        for variant in data['variants']:
            acmg = variant.get('ACMG criteria assessment', '')
            if 'Patogen' in acmg or 'patogen' in acmg:
                finding_status = "PATOGEN VARIANT PÅVISAD"
                break
            elif 'osäker' in acmg.lower() or 'vus' in acmg.lower():
                finding_status = "OKLART FYND"
        
        summary_content.append(finding_status)
        summary_content.append("")  # Empty line
        
        # Identified variants
        summary_content.append(("Identifierade genetiska varianter:", True))
        for i, variant in enumerate(data['variants'], 1):
            gene = variant.get('Gene', 'N/A')
            transcript = variant.get('Transcript', 'N/A')
            nucleotide = variant.get('Nucleotide change', 'N/A')
            protein = variant.get('Protein change', 'N/A')
            zygosity = variant.get('Zygosity', 'N/A')
            
            variant_text = f"Variant {i}: {gene} ({transcript}) c.{nucleotide} p.{protein} - {zygosity}"
            summary_content.append(variant_text)
        
        summary_content.append("")  # Empty line
        
        # Clinical recommendations
        summary_content.append(("Kliniska rekommendationer:", True))
        for variant in data['variants']:
            acmg = variant.get('ACMG criteria assessment', 'N/A')
            clinvar = variant.get('ClinVar and hemophilia database reports', '')
            
            summary_content.append(f"ACMG-klassificering: {acmg}")
            if clinvar:
                summary_content.append(f"Klinisk betydelse: {clinvar}")
        
        summary_content.append("")  # Empty line
        
        # Additional information
        summary_content.append(("Ytterligare information om varianter:", True))
        for variant in data['variants']:
            disease = variant.get('Disease', 'N/A')
            inheritance = variant.get('Inheritance', 'N/A')
            summary_content.append(f"Sjukdom: {disease}")
            summary_content.append(f"Nedärvning: {inheritance}")
        
        summary_content.append("")  # Empty line
        
        # List of analyzed genes
        summary_content.append(("Analyserade gener:", True))
        analyzed_genes = set(v.get('Gene', '') for v in data['variants'])
        summary_content.append(", ".join(sorted(analyzed_genes)))
        
        create_boxed_section(doc, "Sammanfattning och utlåtande:", summary_content)
        
        # === PAGE 1: SAMPLE AND RESULTS BOX ===
        sample_content = []
        
        sample_content.append(("Provtyp:", True))
        sample_content.append("DNA")
        
        sample_content.append("")
        sample_content.append(("Provtagningstidpunkt:", True))
        arrival_date = data.get('Arrival_date', datetime.now().strftime("%Y-%m-%d"))
        sample_content.append(arrival_date)
        
        sample_content.append("")
        sample_content.append(("Externt provnummer:", True))
        sample_content.append(data.get('LID-NR', 'N/A'))
        
        sample_content.append("")
        sample_content.append(("Analysmetod:", True))
        seq_method = data.get('Sequencing method', 'MPS')
        sample_content.append(f"{seq_method} (Massiv parallel sekvensering)")
        
        create_boxed_section(doc, "Prover, analyser och resultat:", sample_content)
        
        # === PAGE 1: FOOTER ===
        create_footer_section(doc)
        
        # === PAGE 2: NEW PAGE ===
        doc.add_page_break()
        
        # === PAGE 2: HEADER ===
        create_header_section(doc, data, page_num=2, total_pages=2)
        
        # === PAGE 2: DETAILED RESULTS TABLE BOX ===
        table_box = doc.add_table(rows=1, cols=1)
        table_box.style = 'Table Grid'
        
        # Apply heavy borders to outer box
        tbl = table_box._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '24')  # Heavy border
            border.set(qn('w:space'), '4')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        cell = table_box.rows[0].cells[0]
        
        # Title
        p = cell.paragraphs[0]
        run = p.add_run("Detaljerade resultat:")
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(6)
        
        # Add methodology description
        p = cell.add_paragraph()
        run = p.add_run("Metodik: WES Genpanel, MPS data")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        p = cell.add_paragraph()
        run = p.add_run("Analysresultat:")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(6)
        
        # Create results table inside the box
        results_table = cell.add_table(rows=1 + len(data['variants']), cols=5)
        results_table.style = 'Table Grid'
        
        # Header row
        headers = ["Gen", "Teoretisk proteinförändring", "Zygositet", "Mutationstyp", "Variantstatus"]
        for i, header in enumerate(headers):
            hcell = results_table.rows[0].cells[i]
            p = hcell.paragraphs[0]
            run = p.add_run(header)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row_idx, variant in enumerate(data['variants'], 1):
            gene = variant.get('Gene', 'N/A')
            protein = f"p.{variant.get('Protein change', 'N/A')}"
            zygosity = variant.get('Zygosity', 'N/A')
            mutation_type = "Missense"  # Could be derived from data
            variant_status = "Patogen" if "Patogen" in variant.get('ACMG criteria assessment', '') else "VUS"
            
            values = [gene, protein, zygosity, mutation_type, variant_status]
            for col_idx, value in enumerate(values):
                vcell = results_table.rows[row_idx].cells[col_idx]
                p = vcell.paragraphs[0]
                run = p.add_run(value)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
        
        # Add interpretation text below table
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run("Tolkning:")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        for variant in data['variants']:
            p = cell.add_paragraph()
            clinvar = variant.get('ClinVar and hemophilia database reports', '')
            if clinvar:
                run = p.add_run(clinvar)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
        
        # === PAGE 2: FOOTER ===
        create_footer_section(doc)

        # Ensure output directory exists
        os.makedirs(output_path, exist_ok=True)

        filename = f"{data['LID-NR']}_{data['variants'][0]['Gene'].upper()}_Svarsrapport.docx"
        output_file = os.path.join(output_path, filename)
        doc.save(output_file)
        logger.info(f"Clinical report successfully created: {output_file}")
        return output_file

    except Exception as e:
        logger.error(f"Error during report generation: {e}", exc_info=True)
        return None
        add_horizontal_line(line_para)
        add_spacer(doc, 6)
        
        sender_info = get_sender_info()
        add_styled_paragraph(doc, sender_info)

        # Ensure output directory exists
        os.makedirs(output_path, exist_ok=True)

        filename = f"{data['LID-NR']}_{data['variants'][0]['Gene'].upper()}.docx"
        output_file = os.path.join(output_path, filename)
        doc.save(output_file)
        logger.info(f"Document successfully created: {output_file}")
        return output_file

    except Exception as e:
        logger.error(f"Error during report generation: {e}", exc_info=True)
        return None
        add_horizontal_line(line_para)
        add_spacer(doc, 6)
        
        sender_info = get_sender_info()
        add_styled_paragraph(doc, sender_info)

        # Ensure output directory exists
        os.makedirs(output_path, exist_ok=True)

        filename = f"{data['LID-NR']}_{data['variants'][0]['Gene'].upper()}.docx"
        output_file = os.path.join(output_path, filename)
        doc.save(output_file)
        logger.info(f"Document successfully created: {output_file}")
        return output_file

    except Exception as e:
        logger.error(f"Error during report generation: {e}", exc_info=True)
        return None

def generate_normalfinding_document(data, output_path):
    """
    Generate a Word document for normal finding (no variants detected) using a template if available.
    
    Args:
        data: Dictionary containing gene and patient information
        output_path: Directory path where the document should be saved
        
    Returns:
        str: Path to the generated document, or None if generation failed
    """
    try:
        # Validate output path
        if not output_path or '..' in output_path:
            logger.error(f"Invalid output path: {output_path}")
            return None
        
        # Try to load template, otherwise create blank document
        template_path = os.path.join(os.path.dirname(__file__), TEMPLATE_FILE)
        if os.path.exists(template_path):
            doc = Document(template_path)
            logger.info(f"Using template for normal finding: {template_path}")
        else:
            doc = Document()
            logger.info("No template found for normal finding, creating blank document")

        gene = data.get("Gene", "okänd gen")
        lid_nr = data.get("LID-NR", "okänt ID")
        transcript = data.get("Transcript", "")
        sequencing_method = data.get("Sequencing method", "okänd metod")

        # Main heading with professional styling
        main_heading = f"Genetisk Rapport: {lid_nr} - {gene.upper()}"
        heading_para = doc.add_paragraph()
        style_heading(heading_para, main_heading, level=1)
        
        # Add decorative line
        line_para = doc.add_paragraph()
        add_horizontal_line(line_para)
        add_spacer(doc, 12)

        # Normal finding section
        add_section_heading(doc, "Normalt Fynd", level=2)
        line_para = doc.add_paragraph()
        add_horizontal_line(line_para)
        add_spacer(doc, 6)
        
        normal_text = config.NORMAL_FINDING_TEXT.format(
            gene=gene,
            transcript=transcript,
            method=sequencing_method
        )
        add_styled_paragraph(doc, normal_text)
        add_spacer(doc, 12)

        # Genomic analysis section
        add_section_heading(doc, "Genomisk Analys", level=2)
        line_para = doc.add_paragraph()
        add_horizontal_line(line_para)
        add_spacer(doc, 6)
        
        genomic_info = process_genomic_analysis_info(data)
        add_styled_paragraph(doc, genomic_info)
        add_spacer(doc, 12)

        # Add references section
        add_references_section(doc)
        add_spacer(doc, 12)

        # Sender information section
        add_section_heading(doc, "Avsändare", level=2)
        line_para = doc.add_paragraph()
        add_horizontal_line(line_para)
        add_spacer(doc, 6)
        
        sender_info = get_sender_info()
        add_styled_paragraph(doc, sender_info)

        # Ensure output directory exists
        os.makedirs(output_path, exist_ok=True)

        filename = f"{lid_nr}_{gene.upper()}_normalfynd.docx"
        output_file = os.path.join(output_path, filename)
        doc.save(output_file)
        logger.info(f"Normal finding document created: {output_file}")
        return output_file
        
    except Exception as e:
        logger.error(f"Error generating normal finding document: {e}", exc_info=True)
        return None

def process_variant_info(variant):
    """
    Process variant information into formatted text.
    
    Args:
        variant: Dictionary containing variant details
        
    Returns:
        str: Formatted variant information
    """
    gene = variant.get('Gene', config.DEFAULT_UNKNOWN_VALUE)
    transcript = variant.get('Transcript', config.DEFAULT_UNKNOWN_VALUE)
    nucleotide_change = variant.get('Nucleotide change', config.DEFAULT_UNKNOWN_VALUE).strip()
    protein_change = variant.get('Protein change', config.DEFAULT_UNKNOWN_VALUE).strip()
    zygosity = variant.get('Zygosity', config.DEFAULT_UNKNOWN_VALUE)
    inheritance = variant.get('Inheritance', config.DEFAULT_UNKNOWN_VALUE)
    disease = variant.get('Disease', config.DEFAULT_UNKNOWN_VALUE)
    
    info = (f"{gene} ({transcript}): c.{nucleotide_change} p.{protein_change} {zygosity}.\n"
            f"Det är troligt att varianten orsakar {disease}.\n"
            f"Nedärvning: {inheritance}.")
    return info


def process_assessment_info(variant):
    """
    Process ACMG assessment and ClinVar information.
    
    Args:
        variant: Dictionary containing assessment details
        
    Returns:
        str: Formatted assessment information
    """
    acmg = variant.get('ACMG criteria assessment', config.DEFAULT_UNKNOWN_VALUE)
    clinvar = variant.get('ClinVar and hemophilia database reports', config.DEFAULT_UNKNOWN_VALUE)
    
    info = (f"Bedömning enligt ACMG-kriterierna: {acmg}.\n"
            f"Sammanfattning och utlåtande: {clinvar}.")
    return info


def process_proband_info(data):
    """
    Process proband information.
    
    Args:
        data: Dictionary containing proband details
        
    Returns:
        str: Formatted proband information
    """
    proband = data.get('Proband', config.DEFAULT_UNKNOWN_VALUE)
    genotype = data.get('Genotype', config.DEFAULT_UNKNOWN_VALUE)
    phenotype = data.get('Phenotype', config.DEFAULT_UNKNOWN_VALUE)

    info = f"{proband}\n"
    if proband != config.DEFAULT_UNKNOWN_VALUE and proband != config.DEFAULT_UNKNOWN_PROBAND:
        info += (f"Genotyp: {genotype}\n"
                 f"Fenotyp: {phenotype}")
    return info


def process_genomic_analysis_info(data):
    """
    Process genomic analysis information based on sequencing method.
    
    Args:
        data: Dictionary containing sequencing method and other details
        
    Returns:
        str: Formatted genomic analysis information
    """
    sequencing_method = data.get('Sequencing method', config.DEFAULT_UNKNOWN_VALUE).strip().lower()
    exon = data.get('Exon', '')
    gene = data.get('Gene', 'genen')

    if sequencing_method == 'sanger':
        return config.GENOMIC_ANALYSIS_SANGER.format(exon=exon, gene=gene)
    else:
        return config.GENOMIC_ANALYSIS_MPS


def get_sender_info():
    """
    Get sender information for the document.
    
    Returns:
        str: Sender information
    """
    return config.SENDER_INFO


def add_references_section(doc):
    """
    Add a professional references section to the document.
    
    Args:
        doc: Document object
    """
    add_section_heading(doc, "Referenser", level=2)
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)
    add_spacer(doc, 6)
    
    references = [
        "Genome Reference Consortium Human Build (2022), Vol. 37.",
        "Clinical genomics (2022), Stockholm, Sverige. Tillgängligt vid: https://www.scilifelab.se/facilities/clinical-genomics-stockholm/",
        "Scout (2022), Tillgängligt vid: https://github.com/Clinical-Genomics/scout",
        "Landrum MJ, Lee JM, Benson M, Brown GR, Chao C, Chitipiralla S, Gu B, Hart J, Hoffman D, Jang W, Karapetyan K, Katz K, Liu C, Maddipatla Z, Malheiro A, McDaniel K, Ovetsky M, Riley G, Zhou G, Holmes JB, Kattman BL, Maglott DR. ClinVar: improving access to variant interpretations and supporting evidence. Nucleic Acids Res, 2018 Jan 4. PubMed PMID: 29165669",
        "EAHAD Coagulation Factor Variant Databases (2022): https://dbs.eahad.org",
        "Richards, Sue et al. Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology. Genetics in medicine: official journal of the American College of Medical Genetics vol. 17,5 (2015): 405-24. doi:10.1038/gim.2015.30"
    ]
    
    for i, ref in enumerate(references, 1):
        para = doc.add_paragraph()
        run = para.add_run(f"{i}. {ref}")
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(3)
